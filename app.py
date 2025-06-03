
import streamlit as st
import pandas as pd
import io, os, tempfile
from PIL import Image, ImageDraw, ImageFont
import qrcode
from barcode import Code39, Code128
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Mm
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from fpdf import FPDF

def mm_to_px(mm): return int(mm * 3.78)
def mm_to_pt(mm): return mm * 2.8346
def mm_to_excel_width(mm): return mm * 0.14

def generate_image(data, code_type, font_size, width_px, height_px):
    if code_type == "QR Code":
        qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_M)
        qr.add_data(data)
        qr.make(fit=True)
        code_img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    elif code_type == "Code 128":
        barcode = Code128(data, writer=ImageWriter())
        output = io.BytesIO()
        barcode.write(output, {"write_text": False})
        output.seek(0)
        code_img = Image.open(output).convert("RGB")
    else:
        barcode = Code39(data, writer=ImageWriter(), add_checksum=False)
        output = io.BytesIO()
        barcode.write(output, {"write_text": False})
        output.seek(0)
        code_img = Image.open(output).convert("RGB")

    try:
        font = ImageFont.truetype("arial.ttf", font_size)
    except:
        font = ImageFont.load_default()

    draw_tmp = ImageDraw.Draw(Image.new("RGB", (1, 1)))
    bbox = draw_tmp.textbbox((0, 0), data, font=font)
    text_width, text_height = bbox[2] - bbox[0], bbox[3] - bbox[1]

    text_img = Image.new("RGB", (width_px, text_height + 10), "white")
    draw = ImageDraw.Draw(text_img)
    draw.text(((width_px - text_width) // 2, 5), data, fill="black", font=font)

    code_height = height_px - text_img.height
    code_img = code_img.resize((width_px, code_height))

    final_img = Image.new("RGB", (width_px, height_px), "white")
    final_img.paste(code_img, (0, 0))
    final_img.paste(text_img, (0, code_height))
    return final_img

def main():
    st.set_page_config(page_title="Générateur stable Render", layout="centered")
    st.title("🧾 Générateur de codes (version 100% Render-compatible)")

    output = None
    filename = ""

    with st.form("formulaire"):
        uploaded_file = st.file_uploader("📄 Fichier Excel", type=["xlsx"])
        format = st.selectbox("📂 Format de sortie", ["PDF", "Word", "Excel"])
        code_type = st.selectbox("🔢 Type de code", ["QR Code", "Code 39", "Code 128"])
        font_size = st.slider("✏️ Taille du texte", 6, 36, 12)
        label_w = st.number_input("📏 Largeur étiquette (mm)", value=50.0)
        label_h = st.number_input("📏 Hauteur étiquette (mm)", value=25.0)
        cols = st.number_input("🧱 Colonnes", min_value=1, max_value=10, value=3)
        rows = st.number_input("🧱 Lignes", min_value=1, max_value=30, value=8)
        spacing_x = st.number_input("↔️ Espacement horizontal (mm)", value=5.0)
        spacing_y = st.number_input("↕️ Espacement vertical (mm)", value=5.0)
        margin_top = st.number_input("⬆️ Marge haute (mm)", value=10.0)
        margin_right = st.number_input("➡️ Marge droite (mm)", value=10.0)
        submitted = st.form_submit_button("Générer")

    if submitted and uploaded_file:
        df = pd.read_excel(uploaded_file)
        codes = df.iloc[:, 0].dropna().astype(str).tolist()
        width_px = mm_to_px(label_w)
        height_px = mm_to_px(label_h)

        if format == "Excel":
            wb = Workbook()
            ws = wb.active
            col_letter = get_column_letter(1)
            ws.column_dimensions[col_letter].width = mm_to_excel_width(label_w)
            for i, code in enumerate(codes, start=1):
                img = generate_image(code, code_type, font_size, width_px, height_px)
                tmp = io.BytesIO()
                img.save(tmp, format='PNG')
                tmp.seek(0)
                xl_img = XLImage(tmp)
                ws.row_dimensions[i].height = mm_to_pt(label_h)
                ws.add_image(xl_img, f"A{i}")
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            filename = "etiquettes.xlsx"

        elif format == "Word":
            doc = Document()
            section = doc.sections[-1]
            section.top_margin = Mm(margin_top)
            section.right_margin = Mm(margin_right)
            section.left_margin = Mm(margin_right)
            table = doc.add_table(rows=int(rows), cols=int(cols))
            table.autofit = False

            i = 0
            for r in range(int(rows)):
                row_cells = table.rows[r].cells
                for c in range(int(cols)):
                    if i < len(codes):
                        code = codes[i]
                        img = generate_image(code, code_type, font_size, width_px, height_px)
                        tmp = io.BytesIO()
                        img.save(tmp, format='PNG')
                        tmp.seek(0)
                        paragraph = row_cells[c].paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(tmp, width=Mm(label_w), height=Mm(label_h))
                        i += 1

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            filename = "etiquettes.docx"

        elif format == "PDF":
            pdf = FPDF(unit="mm", format="A4")
            pdf.add_page()
            page_w, page_h = pdf.w, pdf.h

            x0 = margin_right
            y0 = margin_top
            temp_paths = []

            for i, code in enumerate(codes):
                if i > 0 and i % (cols * rows) == 0:
                    pdf.add_page()

                col = i % cols
                row = (i // cols) % rows

                x = x0 + col * (label_w + spacing_x)
                y = y0 + row * (label_h + spacing_y)

                img = generate_image(code, code_type, font_size, width_px, height_px)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                    img.save(tmp_file.name, format='PNG')
                    temp_paths.append(tmp_file.name)
                    pdf.image(tmp_file.name, x=x, y=y, w=label_w, h=label_h)

            output = io.BytesIO()
            pdf.output(output)
            output.seek(0)
            filename = "etiquettes.pdf"

            for path in temp_paths:
                os.unlink(path)

    if output:
        st.download_button("📥 Télécharger", output, filename)

if __name__ == "__main__":
    main()
